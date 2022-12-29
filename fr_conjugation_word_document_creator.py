import random
from verbecc import Conjugator
cg = Conjugator(lang='fr')




verbes = ["regarder", "entrer", "arriver", "entrer", "rentrer", "monter",
          "rester", "retourner", "tomber", "se lever", "se promener", "espérer", 
          "exagérer", "considérer", "répéter", "préférer", "s'appeler", "jeter",
          "rappeler", "s'ennuyer", "essayer", "nettoyer", "payer", "renvoyer",
          "commencer", "annoncer", "avancer", "effacer", "influencer", "recommencer",
          "remplacer", "manger", "bouger", "changer", "corriger", "déménager", "encourager",
          "s'engager", "interroger", "nager", "obliger", "protéger", "ranger", "télécharger",
          "répondre", "attendre", "défendre", "descendre", "entendre", "vendre", 
          "dormir", "partir", "repartir", "se sentir", "servir", "sortir", "choisir",
          "agir", "applaudir", "définir", "finir", "garantir", "grandir", "réfléchir", "remplir",
          "se réunir", "réussir", "trahir", "craindre", "atteindre", "s'éteindre", "joindre",
          "rejoindre", "se plaindre", "accueillir", "acquérir", "aller", "s'asseoir",
          "avoir", "se battre", "boire", "connaître", "construire", "introduire", "réduire",
          "convaincre", "courir", "croire", "devoir", "dire", "dissoudre", "écrire",
          "envoyer", "être", "faire", "falloir", "fuir", "lire", "mettre", "mourir", 
          "permettre", "promettre", "offrir", "ouvrir", "découvrir", "plaire", "pleuvoir",
          "pouvoir", "prendre", "apprendre", "comprendre", "reprendre", "surprendre",
          "recevoir", "apercevoir", "rire", "sourire", "savoir", "suffire", "suivre",
          "se taire", "venir", "s'abstenir", "appartenir", "devenir", "entretenir", "revenir",
          "obtenir", "prévenir", "se souvenir", "tenir", "vivre", "voir", "vouloir"]
temps = ['présent','passé-composé', 'futur-simple', 'plus-que-parfait', 'imparfait']
solution = []
records = []
for i in range(60):
    temp = temps[random.randint(0,len(temps))-1]
    person = random.randint(0,5)
    verb = verbes[random.randint(0,len(verbes)-1)]
    conjugation = cg.conjugate(verb)
    # print(i+1)
    # print("Verb:", verb)
    # print("Tempus:", temp)
    # print("Numerus:", person+1)
    # print("")
    # print("")
    
    fs = conjugation['moods']['indicatif'][temp][person]
    solution.append((i+1,fs))
    records.append((str(verb)+", "+str(temp)+", "+str(person+1),"_______________________", str(fs)))
    #print(fs)
# print(solution)    





from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Auto-generated french conjugation-training', 0)

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph('first item in unordered list', style='List Bullet')
# document.add_paragraph('first item in ordered list', style='List Number')

#document.add_picture('monty-truth.png', width=Inches(1.25))

# records = [
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# ]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "verb"
hdr_cells[1].text = 'answer'
hdr_cells[2].text = 'solution'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('conjugation-training_fr.docx')
